# easymcat.py
# Streamlit DOCX Study Reader
# - Subject/Topic/Subtopic navigation (markers or Heading 1/2/3)
# - Browser Text-to-Speech (installed voices dropdown)
# - Flow reading: sentence/paragraph pauses (optionally clause pauses)
# - Sticky floating Controls (Listen + Next/Back)
# - Reads mathematical equations (including many Word “Equation” objects) via linearization + speakable math
# - Reads chemical formulas + simple chemical equations more clearly (H2O, NaCl, (NH4)2SO4, CuSO4·5H2O, Fe3+, etc.)
# - Mobile-friendly voice selection (works better on phones/tablets) + prefers Google UK English Male/Female when available
# - Story mode (narration-friendly page formatting)
# - Separate Resume button after Pause
# - Prevent bracketed letters like (A), (B), (C) from being misread as chemical formulas unless truly chemistry-contextual
# - Timeline/progress slider (estimated) with seek + skip forward/back
#
# Run:
#   streamlit run easymcat.py
#
DEFAULT_URL = "https://raw.githubusercontent.com/eogbeide/stock-wizard/main/Exam_Crackers.docx"

import io
import re
import uuid
from typing import Optional, List, Dict, Tuple

import requests
import streamlit as st
from docx import Document

# IMPORTANT: must be first Streamlit command
st.set_page_config(page_title="DOCX Study Reader", layout="wide")

# -----------------------------
# DOCX parsing helpers
# -----------------------------
SUBJECT_RE = re.compile(r"^\s*subject\s*[:\-]\s*(.+?)\s*$", flags=re.IGNORECASE)
TOPIC_RE = re.compile(r"^\s*topic\s*[:\-]\s*(.+?)\s*$", flags=re.IGNORECASE)
SUBTOPIC_RE = re.compile(r"^\s*(?:sub\s*topic|subtopic)\s*[:\-]\s*(.+?)\s*$", flags=re.IGNORECASE)

# WordprocessingML + Office Math namespaces
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def heading_level(style_name: str) -> Optional[int]:
    """Return heading level if style_name looks like 'Heading 1', 'Heading 2', etc."""
    if not style_name:
        return None
    m = re.match(r"Heading\s+(\d+)", str(style_name).strip(), flags=re.IGNORECASE)
    if not m:
        return None
    try:
        return int(m.group(1))
    except Exception:
        return None


def fetch_docx_bytes(url: str) -> bytes:
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    return r.content


def _local(tag: str) -> str:
    return tag.split("}")[-1] if tag else ""


def omml_to_linear(el) -> str:
    """
    Convert common Office Math (OMML) structures into a simple linear text:
      - Fractions -> (num)/(den)
      - Superscripts -> base^sup
      - Subscripts -> base_sub
      - Roots -> sqrt(expr) or root(deg, expr)
    Falls back to concatenating contained text tokens.
    """
    if el is None:
        return ""

    loc = _local(getattr(el, "tag", ""))

    # leaf text in OMML
    if loc == "t":
        return el.text or ""

    # fraction
    if loc == "f":
        num = el.find("m:num", NS)
        den = el.find("m:den", NS)
        num_txt = omml_to_linear(num) if num is not None else ""
        den_txt = omml_to_linear(den) if den is not None else ""
        num_txt = num_txt.strip()
        den_txt = den_txt.strip()
        if num_txt or den_txt:
            return f"({num_txt})/({den_txt})"
        return " ".join(t.text for t in el.findall(".//m:t", NS) if t.text)

    # superscript
    if loc == "sSup":
        base = el.find("m:e", NS)
        sup = el.find("m:sup", NS)
        b = omml_to_linear(base).strip()
        s = omml_to_linear(sup).strip()
        if b or s:
            return f"{b}^{s}"
        return " ".join(t.text for t in el.findall(".//m:t", NS) if t.text)

    # subscript
    if loc == "sSub":
        base = el.find("m:e", NS)
        sub = el.find("m:sub", NS)
        b = omml_to_linear(base).strip()
        s = omml_to_linear(sub).strip()
        if b or s:
            return f"{b}_{s}"
        return " ".join(t.text for t in el.findall(".//m:t", NS) if t.text)

    # root
    if loc == "rad":
        deg = el.find("m:deg", NS)
        expr = el.find("m:e", NS)
        deg_txt = omml_to_linear(deg).strip() if deg is not None else ""
        expr_txt = omml_to_linear(expr).strip() if expr is not None else ""
        if deg_txt:
            return f"root({deg_txt}, {expr_txt})"
        return f"sqrt({expr_txt})"

    # n-ary (sum/integral) - best effort
    if loc == "nary":
        chr_el = el.find("m:chr", NS)
        op = ""
        if chr_el is not None:
            op = chr_el.get(f"{{{NS['m']}}}val") or chr_el.get("m:val") or ""
        sub = el.find("m:sub", NS)
        sup = el.find("m:sup", NS)
        expr = el.find("m:e", NS)
        sub_txt = omml_to_linear(sub).strip() if sub is not None else ""
        sup_txt = omml_to_linear(sup).strip() if sup is not None else ""
        expr_txt = omml_to_linear(expr).strip() if expr is not None else ""
        core = op or "operator"
        if sub_txt:
            core += f"_{sub_txt}"
        if sup_txt:
            core += f"^{sup_txt}"
        if expr_txt:
            core += f"({expr_txt})"
        return core

    # container nodes: recurse children
    parts: List[str] = []
    for c in getattr(el, "iterchildren", lambda: [])():
        chunk = omml_to_linear(c)
        if chunk:
            parts.append(chunk)

    if parts:
        return " ".join(p for p in parts if p).strip()

    toks = [t.text for t in el.findall(".//m:t", NS) if t.text]
    return " ".join(toks).strip()


def paragraph_text_with_math(p) -> str:
    """
    Extract text from a python-docx paragraph including Word equation objects.
    We keep normal runs in order and replace OMML equation blocks with a linear string.
    """
    try:
        pel = p._p  # lxml element
    except Exception:
        return (p.text or "").strip()

    out: List[str] = []

    for child in pel.iterchildren():
        loc = _local(getattr(child, "tag", ""))

        if loc in ("r", "hyperlink", "smartTag", "sdt"):
            for t in child.findall(".//w:t", NS):
                if t.text:
                    out.append(t.text)

        elif loc in ("oMath", "oMathPara"):
            eq = omml_to_linear(child).strip()
            if eq:
                out.append(" ")
                out.append(eq)
                out.append(" ")

        else:
            for t in child.findall(".//w:t", NS):
                if t.text:
                    out.append(t.text)

    joined = "".join(out)
    return joined.strip() if joined.strip() else (p.text or "").strip()


def parse_docx_to_structure(docx_bytes: bytes) -> List[Dict]:
    """
    Structure:
    subjects = [
      {"subject": str, "topics": [{"topic": str, "subtopics": [{"subtopic": str, "chunks": [str], "full_text": str}]}]}
    ]
    """
    doc = Document(io.BytesIO(docx_bytes))

    has_markers = False
    for p in doc.paragraphs:
        raw = paragraph_text_with_math(p)
        if not raw:
            continue
        if SUBJECT_RE.match(raw) or TOPIC_RE.match(raw) or SUBTOPIC_RE.match(raw):
            has_markers = True
            break

    subjects: List[Dict] = []
    cur_subject: Optional[Dict] = None
    cur_topic: Optional[Dict] = None
    cur_subtopic: Optional[Dict] = None

    def ensure_subject(name: str):
        nonlocal cur_subject, cur_topic, cur_subtopic
        subj = {"subject": (name.strip() or "Untitled Subject"), "topics": []}
        subjects.append(subj)
        cur_subject = subj
        cur_topic = None
        cur_subtopic = None

    def ensure_topic(name: str):
        nonlocal cur_subject, cur_topic, cur_subtopic
        if cur_subject is None:
            ensure_subject("General")
        top = {"topic": (name.strip() or "Untitled Topic"), "subtopics": []}
        cur_subject["topics"].append(top)
        cur_topic = top
        cur_subtopic = None

    def ensure_subtopic(name: str):
        nonlocal cur_subject, cur_topic, cur_subtopic
        if cur_subject is None:
            ensure_subject("General")
        if cur_topic is None:
            ensure_topic("Overview")
        sub = {"subtopic": (name.strip() or "Untitled Subtopic"), "chunks": [], "full_text": ""}
        cur_topic["subtopics"].append(sub)
        cur_subtopic = sub

    for p in doc.paragraphs:
        raw = paragraph_text_with_math(p)
        if not raw:
            continue

        if has_markers:
            sm = SUBJECT_RE.match(raw)
            if sm:
                ensure_subject(sm.group(1))
                continue

            tm = TOPIC_RE.match(raw)
            if tm:
                ensure_topic(tm.group(1))
                continue

            stm = SUBTOPIC_RE.match(raw)
            if stm:
                ensure_subtopic(stm.group(1))
                continue

            if cur_topic is not None and cur_subtopic is None:
                ensure_subtopic("Overview")
            if cur_subtopic is not None:
                cur_subtopic["chunks"].append(raw)
            continue

        lvl = heading_level(getattr(p.style, "name", ""))
        if lvl == 1:
            ensure_subject(raw)
            continue
        if lvl == 2:
            ensure_topic(raw)
            continue
        if lvl == 3:
            ensure_subtopic(raw)
            continue

        if cur_subtopic is None:
            ensure_subtopic("Overview")
        cur_subtopic["chunks"].append(raw)

    if not subjects:
        subjects = [
            {
                "subject": "Document",
                "topics": [{"topic": "Content", "subtopics": [{"subtopic": "Overview", "chunks": [], "full_text": ""}]}],
            }
        ]

    for subj in subjects:
        if not subj.get("topics"):
            subj["topics"] = [{"topic": "Overview", "subtopics": [{"subtopic": "Overview", "chunks": [], "full_text": ""}]}]
        for top in subj["topics"]:
            if not top.get("subtopics"):
                top["subtopics"] = [{"subtopic": "Overview", "chunks": [], "full_text": ""}]
            for sub in top["subtopics"]:
                sub["full_text"] = "\n\n".join(sub.get("chunks", [])).strip()

    return subjects


def build_navigation(subjects: List[Dict]) -> Tuple[List[Dict], List[Tuple[int, int, int]]]:
    nav: List[Dict] = []
    flat: List[Tuple[int, int, int]] = []

    for si, subj in enumerate(subjects):
        topics_nav = []
        for ti, top in enumerate(subj.get("topics", [])):
            subs = top.get("subtopics", [])
            if not subs:
                continue
            subs_nav = [{"ui": ui, "subtopic": s.get("subtopic", f"Subtopic {ui+1}")} for ui, s in enumerate(subs)]
            topics_nav.append({"ti": ti, "topic": top.get("topic", f"Topic {ti+1}"), "subtopics": subs_nav})

        if topics_nav:
            nav.append({"si": si, "subject": subj.get("subject", f"Subject {si+1}"), "topics": topics_nav})
            for t in topics_nav:
                for s in t["subtopics"]:
                    flat.append((si, t["ti"], s["ui"]))

    return nav, flat


# -----------------------------
# Page formatting: Story mode
# -----------------------------
def make_story_mode(text: str) -> str:
    """
    Make text more narration-friendly for TTS:
    - Turns bullet/numbered list items into sentences
    - Softens heading lines like 'Key points:' into a spoken lead-in
    - Reduces choppy line breaks (keeps paragraph breaks)
    """
    if not text:
        return ""

    lines = text.splitlines()
    out: List[str] = []

    bullet_re = re.compile(r"^\s*(?:[-•*–—]|\(?\d+\)?[.)])\s+(.*)$")
    heading_re = re.compile(r"^\s*([A-Za-z][A-Za-z0-9 \-/]{2,60})\s*:\s*$")

    for raw in lines:
        line = (raw or "").strip()

        if not line:
            if out and out[-1] != "\n\n":
                out.append("\n\n")
            continue

        hm = heading_re.match(line)
        if hm:
            title = hm.group(1).strip()
            out.append(f"{title}. ")
            continue

        bm = bullet_re.match(line)
        if bm:
            item = (bm.group(1) or "").strip()
            if item:
                if not re.search(r"[.!?]$", item):
                    item += "."
                out.append(item + " ")
            continue

        out.append(line + " ")

    s = "".join(out)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r" *\n\n *", "\n\n", s)
    return s.strip()


# -----------------------------
# Speakable helpers: letters, numbers, math, chemistry
# -----------------------------
_BRACKET_SINGLE_LETTER_RE = re.compile(r"(?<![A-Za-z0-9])[\(\[]\s*([A-Z])\s*[\)\]](?![A-Za-z0-9])")


def make_bracket_letters_speakable(text: str) -> str:
    """
    Convert outline markers like (A), (B), [C] into 'letter A', 'letter B' so they
    are not mistaken for chemical formulas and are spoken clearly.
    """
    if not text:
        return ""

    def repl(m: re.Match) -> str:
        letter = m.group(1)
        return f" letter {letter} "

    return re.sub(_BRACKET_SINGLE_LETTER_RE, repl, text)


_NUM_WORDS_0_19 = {
    0: "zero",
    1: "one",
    2: "two",
    3: "three",
    4: "four",
    5: "five",
    6: "six",
    7: "seven",
    8: "eight",
    9: "nine",
    10: "ten",
    11: "eleven",
    12: "twelve",
    13: "thirteen",
    14: "fourteen",
    15: "fifteen",
    16: "sixteen",
    17: "seventeen",
    18: "eighteen",
    19: "nineteen",
}
_TENS = {20: "twenty", 30: "thirty", 40: "forty", 50: "fifty", 60: "sixty", 70: "seventy", 80: "eighty", 90: "ninety"}


def number_to_words(n: str) -> str:
    """Best-effort small number-to-words for TTS."""
    try:
        val = int(n)
    except Exception:
        return " ".join(list(n))

    if val < 0:
        return "minus " + number_to_words(str(-val))
    if val <= 19:
        return _NUM_WORDS_0_19[val]
    if val < 100:
        tens = (val // 10) * 10
        ones = val % 10
        if ones == 0:
            return _TENS.get(tens, str(val))
        return f"{_TENS.get(tens, str(tens))} {_NUM_WORDS_0_19.get(ones, str(ones))}"
    return " ".join(list(str(val)))


_SUPERSCRIPT_UNICODE = {
    "⁰": "^0",
    "¹": "^1",
    "²": "^2",
    "³": "^3",
    "⁴": "^4",
    "⁵": "^5",
    "⁶": "^6",
    "⁷": "^7",
    "⁸": "^8",
    "⁹": "^9",
}

_SUBSCRIPT_UNICODE = {
    "₀": "0",
    "₁": "1",
    "₂": "2",
    "₃": "3",
    "₄": "4",
    "₅": "5",
    "₆": "6",
    "₇": "7",
    "₈": "8",
    "₉": "9",
}

_SUPER_SIGNS = {"⁺": "+", "⁻": "-", "⁽": "(", "⁾": ")"}

_GREEK = {
    "α": "alpha",
    "β": "beta",
    "γ": "gamma",
    "δ": "delta",
    "ε": "epsilon",
    "θ": "theta",
    "λ": "lambda",
    "μ": "mu",
    "π": "pi",
    "ρ": "rho",
    "σ": "sigma",
    "τ": "tau",
    "φ": "phi",
    "χ": "chi",
    "ψ": "psi",
    "ω": "omega",
}

# Full element symbols + names
ELEMENT_NAMES = {
    "H": "hydrogen",
    "He": "helium",
    "Li": "lithium",
    "Be": "beryllium",
    "B": "boron",
    "C": "carbon",
    "N": "nitrogen",
    "O": "oxygen",
    "F": "fluorine",
    "Ne": "neon",
    "Na": "sodium",
    "Mg": "magnesium",
    "Al": "aluminium",
    "Si": "silicon",
    "P": "phosphorus",
    "S": "sulfur",
    "Cl": "chlorine",
    "Ar": "argon",
    "K": "potassium",
    "Ca": "calcium",
    "Sc": "scandium",
    "Ti": "titanium",
    "V": "vanadium",
    "Cr": "chromium",
    "Mn": "manganese",
    "Fe": "iron",
    "Co": "cobalt",
    "Ni": "nickel",
    "Cu": "copper",
    "Zn": "zinc",
    "Ga": "gallium",
    "Ge": "germanium",
    "As": "arsenic",
    "Se": "selenium",
    "Br": "bromine",
    "Kr": "krypton",
    "Rb": "rubidium",
    "Sr": "strontium",
    "Y": "yttrium",
    "Zr": "zirconium",
    "Nb": "niobium",
    "Mo": "molybdenum",
    "Tc": "technetium",
    "Ru": "ruthenium",
    "Rh": "rhodium",
    "Pd": "palladium",
    "Ag": "silver",
    "Cd": "cadmium",
    "In": "indium",
    "Sn": "tin",
    "Sb": "antimony",
    "Te": "tellurium",
    "I": "iodine",
    "Xe": "xenon",
    "Cs": "caesium",
    "Ba": "barium",
    "La": "lanthanum",
    "Ce": "cerium",
    "Pr": "praseodymium",
    "Nd": "neodymium",
    "Pm": "promethium",
    "Sm": "samarium",
    "Eu": "europium",
    "Gd": "gadolinium",
    "Tb": "terbium",
    "Dy": "dysprosium",
    "Ho": "holmium",
    "Er": "erbium",
    "Tm": "thulium",
    "Yb": "ytterbium",
    "Lu": "lutetium",
    "Hf": "hafnium",
    "Ta": "tantalum",
    "W": "tungsten",
    "Re": "rhenium",
    "Os": "osmium",
    "Ir": "iridium",
    "Pt": "platinum",
    "Au": "gold",
    "Hg": "mercury",
    "Tl": "thallium",
    "Pb": "lead",
    "Bi": "bismuth",
    "Po": "polonium",
    "At": "astatine",
    "Rn": "radon",
    "Fr": "francium",
    "Ra": "radium",
    "Ac": "actinium",
    "Th": "thorium",
    "Pa": "protactinium",
    "U": "uranium",
    "Np": "neptunium",
    "Pu": "plutonium",
    "Am": "americium",
    "Cm": "curium",
    "Bk": "berkelium",
    "Cf": "californium",
    "Es": "einsteinium",
    "Fm": "fermium",
    "Md": "mendelevium",
    "No": "nobelium",
    "Lr": "lawrencium",
    "Rf": "rutherfordium",
    "Db": "dubnium",
    "Sg": "seaborgium",
    "Bh": "bohrium",
    "Hs": "hassium",
    "Mt": "meitnerium",
    "Ds": "darmstadtium",
    "Rg": "roentgenium",
    "Cn": "copernicium",
    "Nh": "nihonium",
    "Fl": "flerovium",
    "Mc": "moscovium",
    "Lv": "livermorium",
    "Ts": "tennessine",
    "Og": "oganesson",
}
ELEMENT_SYMBOLS = set(ELEMENT_NAMES.keys())

COMMON_ACRONYM_BLACKLIST = {"US", "UK", "EU", "UN", "USA", "UAE", "NATO", "MCAT", "SAT", "ACT", "JAMB", "WAEC"}

CHEM_CONTEXT_WORDS = [
    "chemistry",
    "chemical",
    "reaction",
    "reactions",
    "reagent",
    "reagents",
    "equation",
    "equations",
    "acid",
    "acids",
    "base",
    "bases",
    "salt",
    "salts",
    "buffer",
    "buffers",
    "ph",
    "molar",
    "molarity",
    "molality",
    "mole",
    "moles",
    "mol",
    "stoichiometry",
    "stoichiometric",
    "ion",
    "ions",
    "ionic",
    "covalent",
    "oxidation",
    "reduction",
    "redox",
    "electron",
    "electrons",
    "proton",
    "protons",
    "aqueous",
    "precipitate",
    "precipitation",
    "solution",
    "solvent",
    "solute",
    "concentration",
    "dilution",
    "titration",
    "catalyst",
    "catalytic",
    "equilibrium",
    "electrolysis",
    "electrolyte",
    "alkane",
    "alkene",
    "alkyne",
    "alcohol",
    "ester",
    "aldehyde",
    "ketone",
    "amine",
    "benzene",
    "chloride",
    "sulfate",
    "sulphate",
    "nitrate",
    "phosphate",
    "carbonate",
    "bicarbonate",
    "hydroxide",
    "ammonium",
    "sodium",
    "potassium",
    "calcium",
    "magnesium",
    "aluminium",
    "iron",
    "copper",
    "zinc",
    "silver",
    "gold",
    "mercury",
    "lead",
    "bromide",
    "iodide",
    "fluoride",
]
CHEM_CONTEXT_RE = re.compile(r"\b(" + "|".join(map(re.escape, CHEM_CONTEXT_WORDS)) + r")\b", re.IGNORECASE)


def make_math_speakable(text: str, style: str = "Natural") -> str:
    if not text:
        return ""

    s = text

    for k, v in _SUPERSCRIPT_UNICODE.items():
        s = s.replace(k, v)

    for sym, name in _GREEK.items():
        s = s.replace(sym, f" {name} ")

    s = s.replace("≤", " less than or equal to ")
    s = s.replace("≥", " greater than or equal to ")
    s = s.replace("≠", " not equal to ")
    s = s.replace("≈", " approximately equal to ")
    s = s.replace("∝", " proportional to ")
    s = s.replace("∞", " infinity ")
    s = s.replace("→", " tends to ")
    s = s.replace("∑", " summation ")
    s = s.replace("∫", " integral ")
    s = s.replace("√", " square root of ")
    s = s.replace("×", " times ")
    s = s.replace("·", " times ")

    def _root_repl(m: re.Match) -> str:
        deg = m.group(1).strip()
        expr = m.group(2).strip()
        if deg == "2":
            return f" square root of {expr} "
        if deg == "3":
            return f" cube root of {expr} "
        if deg == "4":
            return f" fourth root of {expr} "
        return f" {deg}th root of {expr} "

    s = re.sub(r"\broot\s*\(\s*([0-9]+)\s*,\s*([^)]+)\)", _root_repl, s, flags=re.IGNORECASE)
    s = re.sub(r"\bsqrt\s*\(\s*([^)]+)\)", r" square root of \1 ", s, flags=re.IGNORECASE)

    if style.lower().startswith("lit"):
        s = s.replace("^", " caret ")
        s = s.replace("/", " slash ")
        s = s.replace("=", " equals ")
        s = s.replace("+", " plus ")
        s = s.replace("-", " minus ")
        s = s.replace("*", " times ")
    else:
        s = s.replace("=", " equals ")
        s = s.replace("+", " plus ")
        s = re.sub(r"(?<=\s)-(?=\s)", " minus ", s)
        s = s.replace("*", " times ")

        s = re.sub(r"(\b[A-Za-z][A-Za-z0-9]*)\s*\^\s*2\b", r"\1 squared", s)
        s = re.sub(r"(\b[A-Za-z][A-Za-z0-9]*)\s*\^\s*3\b", r"\1 cubed", s)
        s = re.sub(r"(\b[A-Za-z0-9][A-Za-z0-9]*)\s*\^\s*([A-Za-z0-9]+)\b", r"\1 to the power \2", s)

        s = re.sub(r"(\b[A-Za-z0-9]+)\s*_\s*([A-Za-z0-9]+)\b", r"\1 sub \2", s)
        s = re.sub(r"\(\s*([^)]+?)\s*\)\s*/\s*\(\s*([^)]+?)\s*\)", r" \1 over \2 ", s)
        s = s.replace(" / ", " over ")

    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\s+\n", "\n", s)
    return s.strip()


_FORMULA_RE = re.compile(
    r"""
    (?<![A-Za-z0-9])(
      (?:\d+)?                                     # optional coefficient
      (?:
        (?:\([A-Za-z0-9+\-]+\)\d*)                  # (NH4)2
        |
        (?:[A-Z][a-z]?\d*)                          # Na2, Cl, O2
      )+
      (?:[·\.](?:\d+)?(?:[A-Z][a-z]?\d*)+)*         # ·5H2O  (also accepts .)
      (?:\^?\d*[+\-])?                              # optional charge (Fe3+ or ^2-)
    )(?![A-Za-z0-9])
    """,
    re.VERBOSE,
)


def _normalize_chem_unicode(s: str) -> str:
    if not s:
        return ""
    for k, v in _SUBSCRIPT_UNICODE.items():
        s = s.replace(k, v)
    for k, v in _SUPER_SIGNS.items():
        s = s.replace(k, v)
    # keep math superscripts too (sometimes used for charges)
    for k, v in _SUPERSCRIPT_UNICODE.items():
        s = s.replace(k, v.replace("^", ""))  # ² -> 2 in chemistry context
    s = s.replace("∙", "·")
    return s


def _extract_element_symbols_strict(tok: str) -> Optional[List[str]]:
    """
    Strictly parse element symbols in order.
    Returns None if it contains invalid 'element-looking' tokens.
    """
    t = _normalize_chem_unicode(tok)
    symbols: List[str] = []
    i = 0
    while i < len(t):
        ch = t[i]

        if ch.isspace():
            i += 1
            continue
        if ch.isdigit() or ch in "()[]{}·.+-^":
            i += 1
            continue

        if ch.isalpha():
            if not ch.isupper():
                return None
            sym = ch
            if i + 1 < len(t) and t[i + 1].isalpha() and t[i + 1].islower():
                sym += t[i + 1]
                i += 2
            else:
                i += 1

            if sym not in ELEMENT_SYMBOLS:
                return None
            symbols.append(sym)
            continue

        i += 1

    return symbols if symbols else None


def _is_likely_formula(tok: str, include_single_element: bool, chem_context: bool) -> bool:
    """
    Conservative formula detection that:
      - avoids bracketed letters like (A), (B), (C) being treated as formulas
      - allows ambiguous all-caps short formulas (OH, CO, NO) ONLY in chemistry-context text
    """
    if not tok:
        return False

    t = _normalize_chem_unicode(tok).strip()

    if t in COMMON_ACRONYM_BLACKLIST:
        return False

    # Hard guards for "letters in brackets"
    if re.fullmatch(r"[\(\[]\s*[A-Z]\s*[\)\]]", t):
        return False

    # Roman numerals in parentheses are usually enumeration
    if re.fullmatch(r"\(\s*[IVXLCM]+\s*\)", t) and not re.search(r"[\d]|[+\-^]|[·\.]", t):
        return False

    # Any bracketed pure-letter token with no digits/charge/dot is usually outlining,
    # except in a chemistry context (where (OH) etc can occur).
    m_br = re.fullmatch(r"([\(\[])\s*([A-Za-z]+)\s*([\)\]])", t)
    if m_br and not re.search(r"[\d]|[+\-^]|[·\.]", t):
        inner = m_br.group(2)
        if len(inner) == 1:
            return False
        if not chem_context:
            return False

    symbols = _extract_element_symbols_strict(t)
    if not symbols:
        return False

    has_digit = bool(re.search(r"\d", t))
    has_group = "(" in t or ")" in t
    has_dot = "·" in t or "." in t
    has_charge = bool(re.search(r"(\^?\d*[+\-])$", t))
    has_lower = bool(re.search(r"[a-z]", t))  # indicates 2-letter symbols like Na, Cl
    all_caps_letters_only = bool(re.fullmatch(r"[A-Z]+", t))

    # Prevent plain ALL-CAPS words unless context or stronger hints exist
    if all_caps_letters_only and not (has_digit or has_group or has_dot or has_charge or has_lower or chem_context):
        return False

    # Multi-element formulas:
    if len(symbols) >= 2:
        if has_digit or has_group or has_dot or has_charge or has_lower:
            return True
        return bool(chem_context)

    # Single element:
    if has_digit or has_charge or has_dot:
        return True

    # Allow single-element symbols ONLY when explicitly enabled AND in chemistry context.
    # To avoid false positives in prose (e.g., choice letters A/B/C), only allow 2-letter element symbols
    # like Fe, Na, Cl, Mg when this toggle is on. Single-letter elements still require a digit/charge/dot.
    if include_single_element and chem_context:
        sym = symbols[0]
        if len(sym) >= 2:
            return True

    return False


def _speak_element(sym: str, element_mode: str) -> str:
    if element_mode == "Element names":
        return ELEMENT_NAMES.get(sym, sym)
    return " ".join(list(sym))


def _speak_formula(tok: str, element_mode: str) -> str:
    """
    Turn H2SO4 into: hydrogen two sulfur oxygen four
    Turn (NH4)2SO4 into: open bracket nitrogen hydrogen four close bracket two sulfur oxygen four
    Turn Fe3+ into: iron three plus charge
    """
    t = _normalize_chem_unicode(tok)

    # state symbols
    t = re.sub(r"\(aq\)", " (aq) ", t, flags=re.IGNORECASE)
    t = re.sub(r"\(s\)", " (s) ", t, flags=re.IGNORECASE)
    t = re.sub(r"\(l\)", " (l) ", t, flags=re.IGNORECASE)
    t = re.sub(r"\(g\)", " (g) ", t, flags=re.IGNORECASE)

    out: List[str] = []
    i = 0
    while i < len(t):
        ch = t[i]

        if ch.isspace():
            i += 1
            continue

        if t[i : i + 4].lower() == "(aq)":
            out.append("aqueous")
            i += 4
            continue
        if t[i : i + 3].lower() == "(s)":
            out.append("solid")
            i += 3
            continue
        if t[i : i + 3].lower() == "(l)":
            out.append("liquid")
            i += 3
            continue
        if t[i : i + 3].lower() == "(g)":
            out.append("gas")
            i += 3
            continue

        if ch == "(":
            out.append("open bracket")
            i += 1
            continue
        if ch == ")":
            out.append("close bracket")
            i += 1
            continue
        if ch in ("·", "."):
            out.append("dot")
            i += 1
            continue

        # caret charge format: ^2- or ^-
        if ch == "^":
            j = i + 1
            num = ""
            while j < len(t) and t[j].isdigit():
                num += t[j]
                j += 1
            sign = ""
            if j < len(t) and t[j] in "+-":
                sign = t[j]
                j += 1
            if sign:
                if num:
                    out.append(f"{number_to_words(num)} {'plus' if sign == '+' else 'minus'} charge")
                else:
                    out.append(f"{'plus' if sign == '+' else 'minus'} charge")
            else:
                out.append("charge")
            i = j
            continue

        if ch.isdigit():
            j = i
            num = ""
            while j < len(t) and t[j].isdigit():
                num += t[j]
                j += 1
            out.append(number_to_words(num))
            i = j
            continue

        if ch.isalpha() and ch.isupper():
            sym = ch
            if i + 1 < len(t) and t[i + 1].isalpha() and t[i + 1].islower():
                sym += t[i + 1]
                i += 2
            else:
                i += 1
            out.append(_speak_element(sym, element_mode))
            continue

        if ch in "+-":
            out.append("plus" if ch == "+" else "minus")
            if i == len(t) - 1:
                out.append("charge")
            i += 1
            continue

        out.append(ch)
        i += 1

    spoken = " ".join(out)
    spoken = re.sub(r"[ \t]+", " ", spoken).strip()
    spoken = re.sub(r"\b(plus|minus)\s*$", r"\1 charge", spoken)
    return spoken


def make_chem_speakable(
    text: str,
    element_mode: str = "Element names",
    include_single_element: bool = False,
) -> str:
    """
    Detect likely chemical formulas and replace them with speakable versions.
    Uses chemistry-context heuristics to avoid misreading bracketed letters or outline markers.
    """
    if not text:
        return ""

    s = _normalize_chem_unicode(text)

    # reaction arrows
    s = s.replace("⇌", " reversible reaction ")
    s = s.replace("→", " yields ")
    s = s.replace("⇒", " yields ")
    s = s.replace("<-", " reacts to form ")
    s = s.replace("->", " yields ")

    global_chem_context = bool(CHEM_CONTEXT_RE.search(s))

    def repl(m: re.Match) -> str:
        tok = m.group(1)

        start = m.start(1)
        end = m.end(1)
        window = s[max(0, start - 80) : min(len(s), end + 80)]
        local_context = bool(CHEM_CONTEXT_RE.search(window))
        chem_context = global_chem_context or local_context

        if not _is_likely_formula(tok, include_single_element=include_single_element, chem_context=chem_context):
            return tok

        spoken = _speak_formula(tok, element_mode=element_mode)
        return f" {spoken} "

    s = re.sub(_FORMULA_RE, repl, s)

    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\s+\n", "\n", s)
    return s.strip()


# -----------------------------
# Browser TTS (Web Speech API)
# - Mobile improvements: robust default voice selection & storage by URI OR name/lang
# - Prefer Google UK English Male (en-GB) then Google UK English Female (en-GB) when available
# - Separate Play / Pause / Resume / Stop
# - Estimated timeline slider with seek + skip
# -----------------------------
def tts_component(
    text: str,
    preferred_lang: str = "en-GB",
    rate: float = 0.95,
    pitch: float = 0.78,
    prefer_deep_male_gb: bool = True,
    flow_mode: bool = True,
    sentence_pause_ms: int = 320,
    paragraph_pause_ms: int = 650,
    clause_pause_ms: int = 0,
):
    safe = (text or "").replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${")
    component_id = f"tts_{uuid.uuid4().hex}"

    # Avoid f-strings in the JS body (prevents accidental {var} interpolation errors).
    template = r"""
    <div id="__ID__" style="display:flex; flex-direction:column; gap:10px;">
      <div style="display:flex; gap:10px; align-items:center; flex-wrap:wrap;">
        <button id="__ID___play" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          ▶️ Play
        </button>
        <button id="__ID___pause" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          ⏸ Pause
        </button>
        <button id="__ID___resume" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          🔊 Resume
        </button>
        <button id="__ID___stop" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          ⏹ Stop
        </button>
        <span id="__ID___status" style="color:#666; font-size: 0.9rem;">Ready</span>
      </div>

      <div style="display:flex; gap:10px; align-items:center; flex-wrap:wrap;">
        <button id="__ID___back10" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          ⏪ 10s
        </button>
        <button id="__ID___fwd10" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          10s ⏩
        </button>

        <div style="flex:1; min-width:220px;">
          <input id="__ID___progress" type="range" min="0" max="100" step="0.1" value="0" style="width:100%;"/>
          <div style="display:flex; justify-content:space-between; color:#777; font-size:0.85rem;">
            <span id="__ID___curtime">0:00</span>
            <span id="__ID___totaltime">0:00</span>
          </div>
        </div>
      </div>

      <div style="display:flex; gap:10px; align-items:center; flex-wrap:wrap;">
        <label style="color:#444; font-size:0.9rem;">Voice</label>
        <select id="__ID___voice" style="min-width: 260px; padding:8px 10px; border-radius:10px; border:1px solid #ddd;">
          <option value="">Loading voices…</option>
        </select>
        <button id="__ID___resetvoice" style="padding:8px 12px; border-radius:10px; border:1px solid #ddd; cursor:pointer;">
          Reset default
        </button>
        <span style="color:#888; font-size: 0.85rem;">(Installed voices vary by OS/browser)</span>
      </div>
    </div>

    <script>
      const ROOT_ID = "__ID__";
      const TEXT = `__TEXT__`;

      const playBtn = document.getElementById(ROOT_ID + "_play");
      const pauseBtn = document.getElementById(ROOT_ID + "_pause");
      const resumeBtn = document.getElementById(ROOT_ID + "_resume");
      const stopBtn = document.getElementById(ROOT_ID + "_stop");
      const back10Btn = document.getElementById(ROOT_ID + "_back10");
      const fwd10Btn = document.getElementById(ROOT_ID + "_fwd10");

      const voiceSelect = document.getElementById(ROOT_ID + "_voice");
      const resetVoiceBtn = document.getElementById(ROOT_ID + "_resetvoice");
      const statusEl = document.getElementById(ROOT_ID + "_status");

      const progressEl = document.getElementById(ROOT_ID + "_progress");
      const curTimeEl = document.getElementById(ROOT_ID + "_curtime");
      const totalTimeEl = document.getElementById(ROOT_ID + "_totaltime");

      const preferredLang = "__PREFERRED_LANG__";
      const preferredRate = __RATE__;
      const preferredPitch = __PITCH__;
      const preferDeepMaleGb = __PREFER_DEEP_MALE_GB__;

      const flowMode = __FLOW_MODE__;
      const sentencePauseMs = Math.max(0, __SENTENCE_PAUSE_MS__);
      const paragraphPauseMs = Math.max(0, __PARA_PAUSE_MS__);
      const clausePauseMs = Math.max(0, __CLAUSE_PAUSE_MS__);

      // Bumped key so older saved values don’t override the new default preference.
      // Also helps across mobile browsers where voiceURI can change.
      const storageKey = "easymcat_tts_voice_sel_v4_google_uk_default";

      // Playback state
      // queue items: {text, pauseAfterMs, estSec, cumStartSec, cumEndSec}
      let queue = [];
      let idx = 0;

      let betweenTimer = null;
      let betweenDueAt = 0;
      let betweenRemaining = 0;
      let betweenPaused = false;

      let playing = false;

      // timeline / estimation
      let totalSec = 0;
      let segStartWall = 0;          // ms
      let segStartCumSec = 0;        // seconds
      let segEstSec = 0;             // seconds for current segment
      let uiTicker = null;
      let userDragging = false;

      function setStatus(msg) {
        statusEl.textContent = msg;
      }

      function ensureSpeechSupport() {
        if (!("speechSynthesis" in window) || !("SpeechSynthesisUtterance" in window)) {
          alert("Your browser doesn't support speech synthesis.");
          return false;
        }
        return true;
      }

      function clearBetweenTimer() {
        if (betweenTimer) {
          clearTimeout(betweenTimer);
          betweenTimer = null;
        }
      }

      function fmtTime(sec) {
        sec = Math.max(0, Math.round(sec));
        const m = Math.floor(sec / 60);
        const s = sec % 60;
        return `${m}:${String(s).padStart(2, "0")}`;
      }

      function normalizeText(t) {
        return (t || "")
          .replace(/\r\n/g, "\n")
          .replace(/\r/g, "\n")
          .replace(/[ \t]+/g, " ")
          .replace(/\n[ \t]+/g, "\n")
          .trim();
      }

      function splitParagraphs(t) {
        const cleaned = normalizeText(t);
        if (!cleaned) return [];
        return cleaned
          .split(/\n\s*\n+/g)
          .map(p => p.trim())
          .filter(Boolean);
      }

      function splitSentences(paragraph) {
        const p = (paragraph || "").trim();
        if (!p) return [];

        if (typeof Intl !== "undefined" && Intl.Segmenter) {
          try {
            const seg = new Intl.Segmenter("en", { granularity: "sentence" });
            const parts = [];
            for (const s of seg.segment(p)) {
              const chunk = (s.segment || "").trim();
              if (chunk) parts.push(chunk);
            }
            if (parts.length) return parts;
          } catch (e) {}
        }

        return p
          .split(/(?<=[.!?])\s+/g)
          .map(s => s.trim())
          .filter(Boolean);
      }

      function splitClauses(sentence) {
        const s = (sentence || "").trim();
        if (!s) return [];

        const tokens = s.split(/(,|;|:)/g);
        const out = [];
        let buf = "";

        for (let i = 0; i < tokens.length; i++) {
          const tok = tokens[i];
          if (tok === "," || tok === ";" || tok === ":") {
            buf = (buf + tok).trim();
            if (buf) out.push(buf);
            buf = "";
          } else {
            buf = (buf + " " + tok).trim();
          }
        }
        if (buf.trim()) out.push(buf.trim());
        return out.filter(Boolean);
      }

      // WPM baseline; adjust by rate (roughly linear)
      function estimateSecondsForText(t) {
        const words = (t || "").trim().split(/\s+/).filter(Boolean).length;
        const wpmBase = 155; // comfortable narration
        const wpm = Math.max(60, wpmBase * Math.max(0.5, Math.min(2.0, preferredRate)));
        const sec = (words / wpm) * 60.0;
        return Math.max(0.8, sec); // minimum to keep slider moving
      }

      function buildQueueFromText(t) {
        const paragraphs = splitParagraphs(t);
        const items = [];

        for (let pi = 0; pi < paragraphs.length; pi++) {
          const sentences = splitSentences(paragraphs[pi]);

          for (let si = 0; si < sentences.length; si++) {
            const sent = sentences[si];

            if (flowMode && clausePauseMs > 0) {
              const clauses = splitClauses(sent);
              for (let ci = 0; ci < clauses.length; ci++) {
                const isLastClause = ci === clauses.length - 1;
                const isLastSentence = si === sentences.length - 1;
                const isLastParagraph = pi === paragraphs.length - 1;

                let pauseAfterMs = 0;
                if (!isLastClause) pauseAfterMs = clausePauseMs;
                else if (!isLastSentence) pauseAfterMs = sentencePauseMs;
                else if (!isLastParagraph) pauseAfterMs = paragraphPauseMs;

                const textPart = clauses[ci];
                items.push({ text: textPart, pauseAfterMs, estSec: estimateSecondsForText(textPart) });
              }
            } else if (flowMode) {
              const isLastSentence = si === sentences.length - 1;
              const isLastParagraph = pi === paragraphs.length - 1;

              let pauseAfterMs = 0;
              if (!isLastSentence) pauseAfterMs = sentencePauseMs;
              else if (!isLastParagraph) pauseAfterMs = paragraphPauseMs;

              items.push({ text: sent, pauseAfterMs, estSec: estimateSecondsForText(sent) });
            } else {
              const full = normalizeText(t);
              return [{ text: full, pauseAfterMs: 0, estSec: estimateSecondsForText(full) }];
            }
          }
        }

        if (!items.length) {
          const full = normalizeText(t);
          return [{ text: full, pauseAfterMs: 0, estSec: estimateSecondsForText(full) }];
        }

        return items;
      }

      function computeTimeline(items) {
        let cum = 0;
        for (let i = 0; i < items.length; i++) {
          items[i].cumStartSec = cum;
          cum += (items[i].estSec || 0);
          cum += (Math.max(0, items[i].pauseAfterMs || 0) / 1000.0);
          items[i].cumEndSec = cum;
        }
        totalSec = cum;
        totalTimeEl.textContent = fmtTime(totalSec);
      }

      function updateProgressUI(forceSec) {
        if (userDragging) return;

        let sec = forceSec;
        if (typeof sec !== "number") {
          sec = currentEstimatedSec();
        }
        sec = Math.max(0, Math.min(totalSec || 0, sec));

        if (totalSec > 0) {
          progressEl.value = ((sec / totalSec) * 100.0).toFixed(1);
        } else {
          progressEl.value = "0";
        }
        curTimeEl.textContent = fmtTime(sec);
      }

      function currentEstimatedSec() {
        if (!queue.length) return 0;
        if (!playing) return Math.min(totalSec || 0, (queue[idx]?.cumStartSec || 0));
        const now = Date.now();
        const delta = Math.max(0, (now - segStartWall) / 1000.0);
        const withinSeg = Math.min(segEstSec || 0, delta);
        return Math.min(totalSec || 0, (segStartCumSec || 0) + withinSeg);
      }

      function stopTicker() {
        if (uiTicker) {
          clearInterval(uiTicker);
          uiTicker = null;
        }
      }

      function startTicker() {
        stopTicker();
        uiTicker = setInterval(() => {
          updateProgressUI();
          updateButtons();
        }, 250);
      }

      // -------- Voice handling (mobile-friendly) --------
      function voiceLabel(v) {
        const name = v.name || "Unnamed";
        const lang = v.lang || "";
        return lang ? `${name} (${lang})` : name;
      }

      function voiceKey(v) {
        const uri = v.voiceURI || "";
        const name = v.name || "";
        const lang = v.lang || "";
        if (uri) return "uri:" + uri;
        return "name:" + name + "||" + lang;
      }

      function parseVoiceKey(k) {
        const s = k || "";
        if (s.startsWith("uri:")) return { type: "uri", val: s.slice(4) };
        if (s.startsWith("name:")) {
          const rest = s.slice(5);
          const parts = rest.split("||");
          return { type: "name", name: parts[0] || "", lang: parts[1] || "" };
        }
        return { type: "unknown", raw: s };
      }

      function getVoicesAsync() {
        return new Promise((resolve) => {
          const synth = window.speechSynthesis;
          let voices = synth.getVoices();

          if (voices && voices.length) {
            resolve(voices);
            return;
          }

          const onVoicesChanged = () => {
            voices = synth.getVoices();
            synth.removeEventListener("voiceschanged", onVoicesChanged);
            resolve(voices || []);
          };

          synth.addEventListener("voiceschanged", onVoicesChanged);

          // Mobile sometimes needs a bit longer.
          setTimeout(() => {
            voices = synth.getVoices();
            synth.removeEventListener("voiceschanged", onVoicesChanged);
            resolve(voices || []);
          }, 2000);
        });
      }

      function pickDefaultVoice(voices) {
        if (!voices || !voices.length) return null;

        const norm = (s) => (s || "").toLowerCase();
        const byLangPrefix = (prefix) =>
          voices.filter(v => norm(v.lang).startsWith(prefix.toLowerCase()));

        const gbVoices = byLangPrefix("en-gb");
        const enVoices = byLangPrefix("en");
        const preferredPrefix = preferredLang ? preferredLang.toLowerCase() : "";
        const preferredMatches = preferredPrefix ? byLangPrefix(preferredPrefix) : [];

        // HARD preference: Google UK English Male, then Female (en-GB)
        const googleUkMale = gbVoices.find(v => /google uk english male/i.test(v.name || ""));
        if (googleUkMale) return googleUkMale;

        const googleUkFemale = gbVoices.find(v => /google uk english female/i.test(v.name || ""));
        if (googleUkFemale) return googleUkFemale;

        // Next: deep-male heuristic for GB
        const maleNamePatterns = [
          /google.*uk.*english.*male/i,
          /microsoft.*(ryan|george|alfie).*online/i,
          /microsoft.*(ryan|george|alfie)/i,
          /daniel/i,
          /male/i
        ];

        if (preferDeepMaleGb) {
          for (const re of maleNamePatterns) {
            const found = gbVoices.find(v => re.test(v.name || ""));
            if (found) return found;
          }
          if (gbVoices.length) return gbVoices[0];
        }

        if (preferredMatches.length) return preferredMatches[0];
        if (gbVoices.length) return gbVoices[0];
        if (enVoices.length) return enVoices[0];
        return voices[0];
      }

      function rankVoice(v) {
        const name = (v.name || "").toLowerCase();
        const lang = (v.lang || "").toLowerCase();
        if (lang.startsWith("en-gb") && name.includes("google uk english male")) return 0;
        if (lang.startsWith("en-gb") && name.includes("google uk english female")) return 1;
        return 2;
      }

      function populateVoices(voices) {
        voiceSelect.innerHTML = "";

        const savedKey = (() => {
          try { return localStorage.getItem(storageKey); } catch (e) { return null; }
        })();

        const options = voices.map((v) => ({
          key: voiceKey(v),
          label: voiceLabel(v),
          voice: v
        }));

        options.sort((a, b) => {
          const ra = rankVoice(a.voice), rb = rankVoice(b.voice);
          if (ra !== rb) return ra - rb;
          return a.label.localeCompare(b.label);
        });

        for (const opt of options) {
          const el = document.createElement("option");
          el.value = opt.key;
          el.textContent = opt.label;
          voiceSelect.appendChild(el);
        }

        const hasSaved = savedKey && options.some(o => o.key === savedKey);
        if (hasSaved) {
          voiceSelect.value = savedKey;
          setStatus("Voice loaded (saved).");
          return;
        }

        const def = pickDefaultVoice(voices);
        if (def) {
          voiceSelect.value = voiceKey(def);
          setStatus("Voice loaded (default: Google UK Male if available).");
        } else {
          setStatus("Voice loaded.");
        }
      }

      async function initVoices() {
        if (!ensureSpeechSupport()) {
          setStatus("Speech not supported.");
          return [];
        }
        const voices = await getVoicesAsync();
        if (!voices || !voices.length) {
          voiceSelect.innerHTML = '<option value="">No voices found</option>';
          setStatus("No voices found (device/browser may limit voices).");
          return [];
        }
        populateVoices(voices);
        return voices;
      }

      function getSelectedVoice(voices) {
        const key = voiceSelect.value || "";
        if (!key) return null;
        const parsed = parseVoiceKey(key);

        if (parsed.type === "uri") {
          return voices.find(v => (v.voiceURI || "") === parsed.val) || null;
        }
        if (parsed.type === "name") {
          const n = (parsed.name || "").toLowerCase();
          const l = (parsed.lang || "").toLowerCase();
          return voices.find(v =>
            ((v.name || "").toLowerCase() === n) &&
            ((v.lang || "").toLowerCase() === l)
          ) || voices.find(v => ((v.name || "").toLowerCase() === n)) || null;
        }
        return null;
      }

      // -------- Playback controls --------
      function stopAll() {
        if (!ensureSpeechSupport()) return;

        clearBetweenTimer();
        betweenPaused = false;
        betweenRemaining = 0;

        window.speechSynthesis.cancel();
        playing = false;
        idx = 0;

        setStatus("Stopped.");
        updateButtons();
        updateProgressUI(0);
      }

      function pauseAll() {
        if (!ensureSpeechSupport()) return;

        const synth = window.speechSynthesis;

        // pause while waiting between chunks
        if (betweenTimer) {
          const now = Date.now();
          betweenRemaining = Math.max(0, betweenDueAt - now);
          clearBetweenTimer();
          betweenPaused = true;
          playing = false;
          setStatus("Paused.");
          updateButtons();
          return;
        }

        // pause while speaking
        if (synth.speaking && !synth.paused) {
          synth.pause();
          playing = false;
          setStatus("Paused.");
          updateButtons();
        }
      }

      function resumeAll() {
        if (!ensureSpeechSupport()) return;

        const synth = window.speechSynthesis;

        // resume between chunks
        if (betweenPaused) {
          betweenPaused = false;
          const delay = Math.max(0, betweenRemaining);
          betweenRemaining = 0;
          playing = true;
          scheduleNext(delay);
          setStatus("Speaking…");
          updateButtons();
          return;
        }

        // resume while speaking
        if (synth.paused) {
          synth.resume();
          playing = true;
          setStatus("Speaking…");
          updateButtons();
          return;
        }

        // if not speaking but we have remaining queue
        if (queue.length && idx < queue.length && !synth.speaking) {
          playing = true;
          setStatus("Speaking…");
          speakNext();
          updateButtons();
        }
      }

      function scheduleNext(delayMs) {
        clearBetweenTimer();

        if (delayMs <= 0) {
          speakNext();
          return;
        }

        betweenDueAt = Date.now() + delayMs;
        betweenTimer = setTimeout(() => {
          betweenTimer = null;
          speakNext();
        }, delayMs);
      }

      async function speakItem(itemText, itemEstSec, itemCumStartSec) {
        const voices = await initVoices(); // ensure voices are ready (helps mobile)

        const utter = new SpeechSynthesisUtterance(itemText);
        utter.rate = preferredRate;
        utter.pitch = preferredPitch;
        utter.lang = preferredLang;

        const chosen = getSelectedVoice(voices) || pickDefaultVoice(voices);
        if (chosen) {
          utter.voice = chosen;
          if (chosen.lang) utter.lang = chosen.lang;
        }

        utter.onstart = () => {
          segStartWall = Date.now();
          segStartCumSec = itemCumStartSec || 0;
          segEstSec = itemEstSec || estimateSecondsForText(itemText);

          playing = true;
          setStatus("Speaking…");
          updateButtons();
        };

        utter.onend = () => {
          updateButtons();
        };

        utter.onerror = () => {
          playing = false;
          setStatus("TTS error.");
          updateButtons();
        };

        window.speechSynthesis.speak(utter);
      }

      function speakNext() {
        if (!ensureSpeechSupport()) return;

        const synth = window.speechSynthesis;

        if (idx >= queue.length) {
          playing = false;
          setStatus("Done.");
          updateButtons();
          updateProgressUI(totalSec);
          return;
        }

        const item = queue[idx];
        idx += 1;

        synth.cancel(); // avoid overlap edge-cases
        speakItem(item.text, item.estSec, item.cumStartSec);

        const pauseAfterMs = Math.max(0, item.pauseAfterMs || 0);

        const watcher = setInterval(() => {
          if (!ensureSpeechSupport()) {
            clearInterval(watcher);
            return;
          }
          const s = window.speechSynthesis;

          if (!s.speaking && !s.paused) {
            clearInterval(watcher);
            // segment done; reflect its end time
            updateProgressUI(item.cumEndSec || 0);
            scheduleNext(pauseAfterMs);
          }
        }, 120);
      }

      function startFresh(startSec) {
        if (!ensureSpeechSupport()) return;

        clearBetweenTimer();
        betweenPaused = false;
        betweenRemaining = 0;

        window.speechSynthesis.cancel();
        playing = true;

        queue = buildQueueFromText(TEXT);
        computeTimeline(queue);

        // seek to startSec if provided
        idx = 0;
        if (typeof startSec === "number" && queue.length) {
          const t = Math.max(0, Math.min(totalSec, startSec));
          for (let i = 0; i < queue.length; i++) {
            if (t < (queue[i].cumEndSec || 0)) {
              idx = i;
              break;
            }
          }
        }

        speakNext();
        updateButtons();
        startTicker();
      }

      function play() {
        if (!ensureSpeechSupport()) return;

        const synth = window.speechSynthesis;

        // If paused, Play behaves like Resume
        if (betweenPaused || synth.paused) {
          resumeAll();
          return;
        }

        // If we have remaining queue, continue
        if (queue.length && idx < queue.length && !synth.speaking) {
          playing = true;
          setStatus("Speaking…");
          speakNext();
          updateButtons();
          startTicker();
          return;
        }

        startFresh(0);
      }

      function updateButtons() {
        if (!("speechSynthesis" in window)) {
          playBtn.disabled = false;
          pauseBtn.disabled = true;
          resumeBtn.disabled = true;
          stopBtn.disabled = false;
          return;
        }

        const synth = window.speechSynthesis;
        const speakingNow = synth.speaking && !synth.paused;
        const pausedNow = synth.paused || betweenPaused;
        const hasWork = (queue.length && idx < queue.length) || speakingNow || pausedNow || !!betweenTimer;

        playBtn.disabled = speakingNow;                 // don’t spam Play while speaking
        pauseBtn.disabled = !speakingNow && !betweenTimer;
        resumeBtn.disabled = !pausedNow;
        stopBtn.disabled = !hasWork;
      }

      function seekToSec(t) {
        t = Math.max(0, Math.min(totalSec || 0, t || 0));

        if (!queue.length) {
          // build queue but don't autoplay unless already playing
          queue = buildQueueFromText(TEXT);
          computeTimeline(queue);
        }

        // compute idx
        let newIdx = 0;
        for (let i = 0; i < queue.length; i++) {
          if (t < (queue[i].cumEndSec || 0)) { newIdx = i; break; }
          newIdx = i;
        }
        idx = newIdx;

        // stop current speech and restart from new index
        const synth = window.speechSynthesis;
        clearBetweenTimer();
        betweenPaused = false;
        betweenRemaining = 0;
        synth.cancel();

        // update UI immediately
        updateProgressUI(t);

        // If currently paused, stay paused; otherwise continue speaking
        if (synth.paused) {
          playing = false;
          setStatus("Paused (seeked).");
          updateButtons();
          return;
        }

        playing = true;
        setStatus("Speaking…");
        speakNext();
        updateButtons();
        startTicker();
      }

      // ---- Progress slider wiring ----
      progressEl.addEventListener("input", () => {
        userDragging = true;
        const pct = parseFloat(progressEl.value || "0");
        const t = (pct / 100.0) * (totalSec || 0);
        curTimeEl.textContent = fmtTime(t);
      });

      progressEl.addEventListener("change", () => {
        userDragging = false;
        const pct = parseFloat(progressEl.value || "0");
        const t = (pct / 100.0) * (totalSec || 0);
        seekToSec(t);
      });

      // Skip buttons
      back10Btn.addEventListener("click", () => {
        const t = currentEstimatedSec() - 10;
        seekToSec(t);
      });
      fwd10Btn.addEventListener("click", () => {
        const t = currentEstimatedSec() + 10;
        seekToSec(t);
      });

      // Save selection
      voiceSelect.addEventListener("change", () => {
        try { localStorage.setItem(storageKey, voiceSelect.value || ""); } catch (e) {}
        // If speaking, stop so user can restart with new voice (mobile safer)
        if (ensureSpeechSupport() && (window.speechSynthesis.speaking || window.speechSynthesis.paused)) {
          stopAll();
        } else {
          setStatus("Voice selected.");
          updateButtons();
        }
      });

      resetVoiceBtn.addEventListener("click", async () => {
        try { localStorage.removeItem(storageKey); } catch (e) {}
        const voices = await initVoices();
        const def = pickDefaultVoice(voices);
        if (def) {
          voiceSelect.value = voiceKey(def);
          setStatus("Reset to default voice (Google UK Male if available).");
        } else {
          setStatus("Reset (no default found).");
        }
        if (ensureSpeechSupport() && (window.speechSynthesis.speaking || window.speechSynthesis.paused)) stopAll();
        updateButtons();
      });

      playBtn.addEventListener("click", play);
      pauseBtn.addEventListener("click", pauseAll);
      resumeBtn.addEventListener("click", resumeAll);
      stopBtn.addEventListener("click", stopAll);

      // init on load
      (async function initAll() {
        await initVoices();
        queue = buildQueueFromText(TEXT);
        computeTimeline(queue);
        updateButtons();
        updateProgressUI(0);
        startTicker();
      })();
    </script>
    """

    html = (
        template.replace("__ID__", component_id)
        .replace("__TEXT__", safe)
        .replace("__PREFERRED_LANG__", preferred_lang)
        .replace("__RATE__", str(float(rate)))
        .replace("__PITCH__", str(float(pitch)))
        .replace("__PREFER_DEEP_MALE_GB__", "true" if prefer_deep_male_gb else "false")
        .replace("__FLOW_MODE__", "true" if flow_mode else "false")
        .replace("__SENTENCE_PAUSE_MS__", str(int(sentence_pause_ms)))
        .replace("__PARA_PAUSE_MS__", str(int(paragraph_pause_ms)))
        .replace("__CLAUSE_PAUSE_MS__", str(int(clause_pause_ms)))
    )

    st.components.v1.html(html, height=235)


# -----------------------------
# Streamlit App
# -----------------------------
st.title("DOCX Study Reader")
st.caption("Sidebar: Subject (1) → Topic (2) → Subtopic (3) • Page: content + floating Controls (Listen + Next/Back)")

st.markdown(
    """
<style>
div[data-testid="stHorizontalBlock"] { align-items: flex-start; }
/* Sticky right rail (2nd column) */
div[data-testid="stHorizontalBlock"] > div:nth-child(2) {
  position: sticky;
  top: 4.5rem;
  align-self: flex-start;
  z-index: 5;
}
div[data-testid="stHorizontalBlock"] > div:nth-child(2) > div { padding-top: 0.25rem; }
</style>
""",
    unsafe_allow_html=True,
)


@st.cache_data(show_spinner=True)
def load_structure_from_url(url: str) -> List[Dict]:
    return parse_docx_to_structure(fetch_docx_bytes(url))


with st.sidebar:
    st.header("Document Source")
    url = st.text_input("DOCX URL", value=DEFAULT_URL)
    st.write("Navigation comes from **Subject:** / **Topic:** / **Subtopic:** lines (if present), otherwise Heading 1/2/3.")

try:
    subjects = load_structure_from_url(url)
except Exception as e:
    st.error(f"Could not load DOCX.\n\nError: {e}")
    st.stop()

nav, flat = build_navigation(subjects)
if not nav or not flat:
    st.warning(
        "No usable Subject/Topic/Subtopic sections found.\n\n"
        "Add lines like:\n"
        "- Subject: ...\n"
        "- Topic: ...\n"
        "- Subtopic: ...\n"
        "or use Heading 1/2/3 in the DOCX."
    )
    st.stop()

if "flat_index" not in st.session_state:
    st.session_state.flat_index = 0
st.session_state.flat_index = max(0, min(st.session_state.flat_index, len(flat) - 1))

cur_si, cur_ti, cur_ui = flat[st.session_state.flat_index]

with st.sidebar:
    st.header("Navigate")

    subject_options = [x["subject"] for x in nav]
    cur_subject_nav_idx = next(i for i, x in enumerate(nav) if x["si"] == cur_si)
    selected_subject = st.selectbox("Subject", subject_options, index=cur_subject_nav_idx)

    subj_nav_idx = subject_options.index(selected_subject)
    subj_node = nav[subj_nav_idx]
    new_si = subj_node["si"]

    topic_options = [t["topic"] for t in subj_node["topics"]]
    if new_si == cur_si:
        topic_default_idx = next((i for i, t in enumerate(subj_node["topics"]) if t["ti"] == cur_ti), 0)
    else:
        topic_default_idx = 0
    selected_topic = st.selectbox("Topic", topic_options, index=topic_default_idx)

    topic_nav_idx = topic_options.index(selected_topic)
    topic_node = subj_node["topics"][topic_nav_idx]
    new_ti = topic_node["ti"]

    subtopic_options = [s["subtopic"] for s in topic_node["subtopics"]]
    if new_si == cur_si and new_ti == cur_ti:
        subtopic_default_idx = next((i for i, s in enumerate(topic_node["subtopics"]) if s["ui"] == cur_ui), 0)
    else:
        subtopic_default_idx = 0
    selected_subtopic = st.selectbox("Subtopic", subtopic_options, index=subtopic_default_idx)

    sub_nav_idx = subtopic_options.index(selected_subtopic)
    new_ui = topic_node["subtopics"][sub_nav_idx]["ui"]

    if st.button("Go", use_container_width=True):
        for idx, (si, ti, ui) in enumerate(flat):
            if si == new_si and ti == new_ti and ui == new_ui:
                st.session_state.flat_index = idx
                break
        st.rerun()

    st.divider()
    st.subheader("Text-to-Speech")

    preferred_lang = st.selectbox(
        "Preferred language",
        ["en-GB", "en-US", "en", "es-ES", "fr-FR"],
        index=0,
        help="Voice dropdown uses your installed voices. This sets the language preference/fallback.",
    )

    rate = st.slider("Rate", 0.5, 2.0, 0.95, 0.05)
    pitch = st.slider("Pitch", 0.5, 2.0, 0.78, 0.05, help="Lower pitch generally sounds deeper.")
    prefer_deep_male_gb = st.toggle(
        "Prefer deep male UK voice (default)",
        value=True,
        help="Also prioritises Google UK English Male/Female (en-GB) when available on device.",
    )

    st.divider()
    st.subheader("Flow (pauses)")

    flow_mode = st.toggle(
        "Flow mode (sentence pacing)",
        value=True,
        help="Speaks sentence-by-sentence with short pauses, for a more 'coach-like' delivery.",
    )
    sentence_pause_ms = st.slider("Pause after sentence (ms)", 0, 1500, 320, 20)
    paragraph_pause_ms = st.slider("Pause after paragraph (ms)", 0, 4000, 650, 50)
    clause_pause_ms = st.slider(
        "Optional pause after clause (ms)",
        0,
        800,
        0,
        10,
        help="If > 0, splits sentences by commas/;/: and adds a light pause for clearer rhythm.",
    )

    st.divider()
    st.subheader("Page formatting")

    story_mode = st.toggle(
        "Make it narration-friendly (story mode)",
        value=False,
        help="Smooths bullets/line breaks into more natural narration for TTS.",
    )

    st.divider()
    st.subheader("Chemistry reading")

    read_chem = st.toggle(
        "Speak chemical formulas clearly",
        value=True,
        help="Turns formulas like H2O / NaCl / (NH4)2SO4 / CuSO4·5H2O into clearer spoken text. Avoids bracketed letters like (A) being misread.",
    )
    chem_element_mode = st.selectbox(
        "Chemical element style",
        ["Element names", "Spell symbols"],
        index=0,
        disabled=not read_chem,
        help="Element names: 'sodium chloride'. Spell symbols: 'N A C L'.",
    )
    include_single_element = st.toggle(
        "Include single-element symbols (e.g., Fe)",
        value=False,
        disabled=not read_chem,
        help="Off by default to reduce false positives; when enabled, only triggers in chemistry-context passages.",
    )

    st.divider()
    st.subheader("Math reading")

    read_math = st.toggle(
        "Speak equations clearly",
        value=True,
        help="Converts math symbols to words and extracts many Word equation objects.",
    )
    math_style = st.selectbox("Math style", ["Natural", "Literal"], index=0, disabled=not read_math)

    st.divider()
    show_tts_preview = st.toggle("Show TTS text preview", value=False)

# Current content
cur_si, cur_ti, cur_ui = flat[st.session_state.flat_index]
cur_subject = subjects[cur_si]["subject"]
cur_topic = subjects[cur_si]["topics"][cur_ti]["topic"]
cur_subtopic = subjects[cur_si]["topics"][cur_ti]["subtopics"][cur_ui]["subtopic"]
cur_text = (subjects[cur_si]["topics"][cur_ti]["subtopics"][cur_ui].get("full_text") or "").strip()

tts_text = cur_text

if tts_text:
    # Always convert outline letters like (A) into "letter A" for speech clarity.
    tts_text = make_bracket_letters_speakable(tts_text)

if tts_text and story_mode:
    tts_text = make_story_mode(tts_text)

if tts_text and read_chem:
    tts_text = make_chem_speakable(
        tts_text,
        element_mode=chem_element_mode,
        include_single_element=include_single_element,
    )

if tts_text and read_math:
    tts_text = make_math_speakable(tts_text, style=math_style)

col_left, col_right = st.columns([2, 1], vertical_alignment="top")

with col_left:
    st.subheader(f"{cur_subject}  →  {cur_topic}  →  {cur_subtopic}")

    if cur_text:
        st.write(cur_text)
        if show_tts_preview and tts_text and tts_text != cur_text:
            with st.expander("TTS preview (what will be spoken)"):
                st.write(tts_text)
    else:
        st.info("No paragraph text under this subtopic.")

with col_right:
    st.subheader("Controls")

    st.markdown("**Listen**")
    if tts_text:
        tts_component(
            tts_text,
            preferred_lang=preferred_lang,
            rate=rate,
            pitch=pitch,
            prefer_deep_male_gb=prefer_deep_male_gb,
            flow_mode=flow_mode,
            sentence_pause_ms=sentence_pause_ms,
            paragraph_pause_ms=paragraph_pause_ms,
            clause_pause_ms=clause_pause_ms,
        )
    else:
        st.caption("Nothing to read for this subtopic.")

    st.divider()
    st.markdown("**Navigate**")

    if st.button("⬅️ Back", disabled=(st.session_state.flat_index == 0), use_container_width=True):
        st.session_state.flat_index -= 1
        st.rerun()

    if st.button("Next ➡️", disabled=(st.session_state.flat_index == len(flat) - 1), use_container_width=True):
        st.session_state.flat_index += 1
        st.rerun()

    st.divider()
    st.caption("Progress")
    st.write(f"Section {st.session_state.flat_index + 1} of {len(flat)}")
