import streamlit as st
import requests
import io
import re
from urllib.parse import urlsplit, urlunsplit, quote
from docx import Document  # pip install python-docx

# =========================
# Config
# =========================
st.set_page_config(page_title="Psych 180 MCQs", page_icon="🧠", layout="centered")
DEFAULT_DOC_URL = "https://raw.githubusercontent.com/eogbeide/stock-wizard/main/Psych 180 Pages.docx"

# Mobile-friendly tweaks
st.markdown("""
<style>
button, .stButton>button {padding: 0.65rem 1.1rem; font-size: 1rem;}
div[role="radiogroup"] label {padding: 0.25rem 0; line-height: 1.6;}
</style>
""", unsafe_allow_html=True)

# =========================
# Helpers
# =========================
def encode_url(url: str) -> str:
    parts = urlsplit(url)
    return urlunsplit((parts.scheme, parts.netloc, quote(parts.path), parts.query, parts.fragment))

def _norm(s: str) -> str:
    return re.sub(r'\s+', ' ', str(s or "").strip())

def is_question_start(line: str) -> bool:
    return bool(re.match(r'^\s*(?:Q(?:uestion)?\s*)?\d+\s*[\.\)]\s*\S', line, flags=re.I))

def strip_qnum(line: str) -> str:
    return re.sub(r'^\s*(?:Q(?:uestion)?\s*)?\d+\s*[\.\)]\s*', '', line, flags=re.I)

def split_meta_chain(text: str):
    """
    Extract Answer / Explanation / Rationale tokens even when chained on one line.
    Returns dict with 'answer','explanation','rationale','leftover' (non-meta).
    Accepts 'Rational' (common typo) and 'Why/Exp' for explanation.
    """
    out = {"answer": None, "explanation": None, "rationale": None, "leftover": ""}
    t = text or ""
    pat = re.compile(r'(?i)(Answer|Ans|Key|Correct|Explanation|Why|Exp|Rationale|Rational)\s*:\s*')
    matches = list(pat.finditer(t))
    if not matches:
        # treat "Eliminate ..." lines as rationale if present
        tt = _norm(t)
        if re.match(r'^(Eliminate|Because)\b', tt, flags=re.I):
            out["rationale"] = tt
            out["leftover"] = ""
        else:
            out["leftover"] = tt
        return out

    # leftover before first tag = stem text
    before = t[:matches[0].start()].strip()
    if before:
        out["leftover"] = _norm(before)

    def put(label, value):
        if not value:
            return
        if label in ("answer", "ans", "key", "correct"):
            out["answer"] = _norm(value)
        elif label in ("explanation", "why", "exp"):
            out["explanation"] = _norm(value) if not out["explanation"] else _norm(out["explanation"] + " " + value)
        elif label in ("rationale", "rational"):
            out["rationale"] = _norm(value) if not out["rationale"] else _norm(out["rationale"] + " " + value)

    for i, m in enumerate(matches):
        lbl = m.group(1).lower()
        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(t)
        seg = t[start:end].strip()
        put(lbl, seg)

    return out

def parse_inline_options(line: str):
    """
    Parse a line that may contain inline options like:
      "A. text B) text C: text D - text F. Answer: C Explanation: ..."
    Returns (leading_text, parts) where parts is [('A','...'), ('B','...'), ...]
    """
    parts = []
    pat = re.compile(r'\(?([A-Fa-f])\)?\s*[\.\):\-]\s*')
    matches = list(pat.finditer(line))
    if not matches:
        return (_norm(line), [])

    lead = _norm(line[:matches[0].start()])
    for i, m in enumerate(matches):
        letter = m.group(1).upper()
        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(line)
        seg = _norm(line[start:end])
        if seg:
            parts.append((letter, seg))
    return (lead, parts)

def parse_single_option(line: str):
    m = re.match(r'^\s*\(?([A-Fa-f])\)?\s*[\.\):\-]\s*(\S.*)$', line)
    if m:
        return (m.group(1).upper(), _norm(m.group(2)))
    return None

def figure_out_correct_idx(correct_val, letters, texts):
    if correct_val is None:
        return 0
    cv = _norm(correct_val)

    # Letter key?
    if re.fullmatch(r'[A-Ea-e]', cv):
        L = cv.upper()
        if L in letters:
            return letters.index(L)

    # Exact text match
    for i, t in enumerate(texts):
        if _norm(t).lower() == cv.lower():
            return i

    # Contains match
    for i, t in enumerate(texts):
        if cv.lower() in _norm(t).lower():
            return i

    return 0

def finalize_letters_texts(collected):
    """
    Keep only A–D in A..D order (drop E/F etc.).
    """
    real = [(L, T) for (L, T) in collected if L in ('A','B','C','D')]
    letters = [L for L, _ in real][:4]
    texts   = [T for _, T in real][:4]
    return letters, texts

# =========================
# Load & Parse DOCX
# =========================
@st.cache_data(show_spinner=True)
def load_questions(doc_url: str):
    resp = requests.get(encode_url(doc_url), timeout=30)
    resp.raise_for_status()
    doc = Document(io.BytesIO(resp.content))

    # Lines from paragraphs (fallback to tables)
    lines = []
    for p in doc.paragraphs:
        t = _norm(p.text)
        if t or (lines and lines[-1] != ""):
            lines.append(t if t else "")
    if not any(lines):
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = _norm(p.text)
                        if t or (lines and lines[-1] != ""):
                            lines.append(t if t else "")

    questions = []
    q_text_parts = []
    option_pairs = []          # [('A','...'), ('B','...'), ...]
    correct_val = None
    explanation_parts = []
    rationale_parts = []
    collecting_expl = False
    collecting_rat = False

    def flush():
        if q_text_parts and option_pairs:
            letters, texts = finalize_letters_texts(option_pairs)
            if not letters or not texts:
                return
            q_text = _norm(" ".join(q_text_parts))
            # strip "Question:" if present
            q_text = re.sub(r'^\s*(?:Question)\s*[:\-]\s*', '', q_text, flags=re.I)
            # defensively drop any stray meta that slipped into stem
            q_text = re.sub(r'\b(?:Answer|Ans|Key|Correct|Explanation|Why|Exp|Rationale|Rational)\s*:\s*.*$', '', q_text, flags=re.I)
            correct_idx = figure_out_correct_idx(correct_val, letters, texts)
            questions.append({
                "question": q_text,
                "letters": letters,
                "options": texts,
                "correct_idx": max(0, min(correct_idx, len(texts)-1)),
                "explanation": _norm(" ".join(explanation_parts)) if explanation_parts else "",
                "rationale": _norm(" ".join(rationale_parts)) if rationale_parts else ""
            })

    for raw in lines:
        line = raw.strip()

        # Collecting Explanation
        if collecting_expl:
            if is_question_start(line):
                flush()
                q_text_parts, option_pairs = [strip_qnum(line)], []
                correct_val, explanation_parts, rationale_parts = None, [], []
                collecting_expl = collecting_rat = False
                continue
            meta = split_meta_chain(line)
            if meta["rationale"]:
                collecting_expl = False
                collecting_rat = True
                rationale_parts.append(meta["rationale"])
            if meta["explanation"]:
                explanation_parts.append(meta["explanation"])
            # ignore meta leftovers while in explanation
            continue

        # Collecting Rationale
        if collecting_rat:
            if is_question_start(line):
                flush()
                q_text_parts, option_pairs = [strip_qnum(line)], []
                correct_val, explanation_parts, rationale_parts = None, [], []
                collecting_expl = collecting_rat = False
                continue
            meta = split_meta_chain(line)
            if meta["rationale"]:
                rationale_parts.append(meta["rationale"])
            # ignore leftovers while in rationale
            continue

        # New question?
        if is_question_start(line):
            if q_text_parts or option_pairs:
                flush()
                q_text_parts, option_pairs = [], []
                correct_val, explanation_parts, rationale_parts = None, [], []
            q_text_parts = [strip_qnum(line)]
            continue

        # Inline options?
        lead, parts = parse_inline_options(line)
        if lead:
            meta = split_meta_chain(lead)
            if meta["answer"]:
                correct_val = meta["answer"]
            if meta["explanation"]:
                collecting_expl = True
                explanation_parts.append(meta["explanation"])
            if meta["rationale"]:
                collecting_rat = True
                rationale_parts.append(meta["rationale"])
            if meta["leftover"]:
                q_text_parts.append(meta["leftover"])

        if parts:
            for L, seg in parts:
                meta = split_meta_chain(seg)
                if meta["answer"]:
                    correct_val = meta["answer"]
                if meta["explanation"]:
                    collecting_expl = True
                    explanation_parts.append(meta["explanation"])
                if meta["rationale"]:
                    collecting_rat = True
                    rationale_parts.append(meta["rationale"])
                if meta["leftover"] and L in ('A','B','C','D'):
                    option_pairs.append((L, meta["leftover"]))
            continue

        # Single option on its own line?
        opt = parse_single_option(line)
        if opt:
            L, seg = opt
            meta = split_meta_chain(seg)
            if meta["answer"]:
                correct_val = meta["answer"]
            if meta["explanation"]:
                collecting_expl = True
                explanation_parts.append(meta["explanation"])
            if meta["rationale"]:
                collecting_rat = True
                rationale_parts.append(meta["rationale"])
            if meta["leftover"] and L in ('A','B','C','D'):
                option_pairs.append((L, meta["leftover"]))
            continue

        # Pure meta line (no A/B/C prefix)
        meta = split_meta_chain(line)
        if meta["answer"] or meta["explanation"] or meta["rationale"]:
            if meta["answer"]:
                correct_val = meta["answer"]
            if meta["explanation"]:
                collecting_expl = True
                explanation_parts.append(meta["explanation"])
            if meta["rationale"]:
                collecting_rat = True
                rationale_parts.append(meta["rationale"])
            # ignore leftover in meta lines
            continue

        # Empty line → soft separator
        if line == "":
            continue

        # Otherwise, stem continuation
        q_text_parts.append(line)

    # Flush last block
    flush()

    if not questions:
        raise ValueError(
            "No usable questions parsed. Use numbered questions (e.g., '1.'), A–D options, and an 'Answer:' line. "
            "Optional 'Explanation:' and 'Rationale:' (or 'Rational') supported."
        )
    return questions

# =========================
# App UI
# =========================
st.title("🧠 Psych 180 MCQs (from .docx)")
doc_url = st.text_input("DOCX URL", DEFAULT_DOC_URL)
if not doc_url:
    st.stop()

try:
    qs = load_questions(doc_url)
except Exception as e:
    st.error(f"Failed to load/parse DOCX: {e}")
    st.stop()

# Session state
if "q_idx" not in st.session_state:
    st.session_state.q_idx = 0
if "selection" not in st.session_state:
    st.session_state.selection = {}       # per-question selected VALUE: -1 or 0..3
if "submitted" not in st.session_state:
    st.session_state.submitted = {}       # per-question submitted flag
if "checked_value" not in st.session_state:
    st.session_state.checked_value = {}   # per-question last submitted VALUE

total = len(qs)
st.caption(f"Loaded {total} question(s) from Psych 180 Pages")

# Navigation
c1, c2, c3 = st.columns([1,2,1])
with c1:
    st.button("◀ Back", disabled=(st.session_state.q_idx == 0),
              on_click=lambda: st.session_state.update(q_idx=max(0, st.session_state.q_idx - 1)))
with c3:
    st.button("Next ▶", disabled=(st.session_state.q_idx >= total - 1),
              on_click=lambda: st.session_state.update(q_idx=min(total - 1, st.session_state.q_idx + 1)))

idx = st.session_state.q_idx
q = qs[idx]

# Question text (clean; no meta mixed in)
st.subheader(f"Q{idx+1}. {q['question']}")

# Build radio list with A–D
labels = [f"{q['letters'][i]}. {q['options'][i]}" for i in range(len(q['options']))]

# Radio with sentinel until user chooses
SENTINEL = -1
values = [SENTINEL] + list(range(len(labels)))  # [-1, 0, 1, 2, 3]
prev = st.session_state.selection.get(idx, SENTINEL)
try:
    default_index = values.index(prev)
except ValueError:
    default_index = 0

selected_value = st.radio(
    "Choose one:",
    options=values,
    format_func=lambda v: "— Select an answer —" if v == SENTINEL else labels[v],
    index=default_index,
    key=f"radio_{idx}"
)

# Persist current choice
st.session_state.selection[idx] = selected_value

# If user changes selection after submit, require re-submit
if st.session_state.submitted.get(idx, False):
    if st.session_state.checked_value.get(idx, SENTINEL) != selected_value:
        st.session_state.submitted[idx] = False
        st.info("Selection changed — click Submit to check.")

# Submit
if st.button("Submit", key=f"submit_{idx}"):
    st.session_state.submitted[idx] = True
    st.session_state.checked_value[idx] = st.session_state.selection.get(idx, SENTINEL)

# Result after submission
if st.session_state.submitted.get(idx, False):
    checked = st.session_state.checked_value.get(idx, SENTINEL)
    if checked == SENTINEL:
        st.warning("Please select an option before submitting.")
    else:
        # Compare selection to correct_idx
        is_correct = (checked == q["correct_idx"])
        if is_correct:
            st.success("✅ Correct!")
        else:
            st.error("❌ Incorrect — try again.")

        # Reveal ANSWER and RATIONALE (as requested) at the bottom
        ans_letter = q["letters"][q["correct_idx"]]
        ans_text   = q["options"][q["correct_idx"]]
        st.markdown(f"**Answer:** {ans_letter}. {ans_text}")
        if q.get("rationale"):
            st.markdown(f"**Rationale:** {q['rationale']}")
        # (If you also want to show explanation, uncomment below)
        # if q.get('explanation'):
        #     st.markdown(f"**Explanation:** {q['explanation']}")

with st.expander("Jump to question"):
    jump = st.slider("Question number", 1, total, idx+1)
    if jump - 1 != idx:
        st.session_state.q_idx = jump - 1
        st.rerun()
